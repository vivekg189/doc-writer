"""Document generation routes."""
from flask import render_template, request, flash, session, jsonify
from . import document_bp
from app.services.processor import LegalDocumentProcessor
from app.models.history import add_user_history, save_generated_document
import spacy
import json
import tempfile
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.units import inch
from datetime import datetime
try:
    from dateutil.relativedelta import relativedelta
except ImportError:
    # Fallback if dateutil is not available
    relativedelta = None
import re
import requests


nlp = spacy.load('en_core_web_sm')
processor = LegalDocumentProcessor()

def translate_text(text, target_lang):
    """Translate text using MyMemory API"""
    if not text.strip():
        return text

    try:
        url = f"https://api.mymemory.translated.net/get?q={text}&langpair=en|{target_lang}"
        print(f"Translating: '{text}' to {target_lang} using URL: {url}")
        response = requests.get(url, timeout=5)
        data = response.json()
        print(f"Translation response: {data}")
        if data['responseStatus'] == 200:
            translated = data['responseData']['translatedText']
            print(f"Translated '{text}' to '{translated}'")
            return translated
        else:
            print(f"Translation failed for '{text}'")
            return text  # Fallback to original text
    except Exception as e:
        print(f"Translation error for '{text}': {e}")
        return text  # Fallback to original text

def get_default_data_for_document(doc_type, language):
    """Get default data for realistic document generation"""
    current_date = datetime.now()
    
    defaults = {
        'date': current_date.strftime('%B %d, %Y'),
        'month': current_date.strftime('%B'),
        'year': current_date.strftime('%Y'),
        'execution_date': current_date.strftime('%d/%m/%Y'),
        'execution_place': 'Chennai',
        'witness1_name': 'Mr. Witness One',
        'witness1_address': 'No. 123, Main Street, Chennai - 600001',
        'witness2_name': 'Mr. Witness Two', 
        'witness2_address': 'No. 456, Second Street, Chennai - 600002',
        'jurisdiction': 'Chennai',
        'registration_office': 'Chennai',
        'stamp_duty_bearer': 'Vendee'
    }
    
    if doc_type == 'rental_agreement':
        defaults.update({
            'owner_age': '45',
            'renter_age': '35',
            'owner_father': 'Father Name',
            'renter_father': 'Father Name',
            'owner_address': 'No. 123, Main Street',
            'owner_city': 'Chennai',
            'owner_pincode': '600001',
            'renter_address': 'No. 456, Second Street',
            'renter_city': 'Chennai',
            'renter_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'start_date': '1st April 2024',
            'effective_date': '1st April 2024',
            'duration': '11',
            'renewal_period': '11',
            'rent_amount': '15,000',
            'rent_amount_words': 'Fifteen Thousand',
            'rent_due_date': '1st',
            'rent_increase_percentage': '10',
            'security_deposit': '30,000',
            'security_deposit_words': 'Thirty Thousand',
            'notice_period': '2'
        })
    elif doc_type == 'land_sale_deed':
        defaults.update({
            'seller_age': '50',
            'buyer_age': '40',
            'seller_father': 'Father Name',
            'buyer_father': 'Father Name',
            'seller_address': 'No. 123, Main Street',
            'seller_city': 'Chennai',
            'seller_pincode': '600001',
            'buyer_address': 'No. 456, Second Street',
            'buyer_city': 'Chennai',
            'buyer_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'sale_amount': '50,00,000',
            'sale_amount_words': 'Fifty Lakhs',
            'sale_date': '1st April 2024',
            'survey_number': '123/45',
            'area': '2400',
            'north_boundary': 'Main Road',
            'south_boundary': 'Residential Area',
            'east_boundary': 'Park',
            'west_boundary': 'Commercial Area'
        })
    elif doc_type == 'power_of_attorney':
        defaults.update({
            'principal_age': '55',
            'attorney_age': '40',
            'principal_father': 'Father Name',
            'attorney_father': 'Father Name',
            'principal_address': 'No. 123, Main Street',
            'principal_city': 'Chennai',
            'principal_pincode': '600001',
            'attorney_address': 'No. 456, Second Street',
            'attorney_city': 'Chennai',
            'attorney_pincode': '600002',
            'matter_description': 'property management and legal representation',
            'effective_date': '1st April 2024',
            'expiry_date': '31st March 2025'
        })
    elif doc_type == 'house_lease':
        # Calculate end date based on lease period
        def calculate_end_date(start_date_str, lease_period_years):
            try:
                if relativedelta:
                    # Handle different date formats
                    for fmt in ['%dst %B %Y', '%dnd %B %Y', '%drd %B %Y', '%dth %B %Y', '%d %B %Y']:
                        try:
                            start_date_obj = datetime.strptime(start_date_str, fmt)
                            break
                        except ValueError:
                            continue
                    else:
                        # Fallback to current date
                        start_date_obj = datetime.now()
                    
                    end_date_obj = start_date_obj + relativedelta(years=int(lease_period_years))
                    day = end_date_obj.day
                    if day == 1:
                        suffix = 'st'
                    elif day == 2:
                        suffix = 'nd'
                    elif day == 3:
                        suffix = 'rd'
                    else:
                        suffix = 'th'
                    return end_date_obj.strftime(f'%d{suffix} %B %Y')
                else:
                    # Simple fallback calculation without dateutil
                    return '31st March 2026'
            except:
                return '31st March 2026'
        
        end_date_str = calculate_end_date('1st April 2024', '2')
            
        defaults.update({
            'lessor_age': '50',
            'lessee_age': '35',
            'lessor_father': 'Father Name',
            'lessee_father': 'Father Name',
            'lessor_address': 'No. 123, Main Street',
            'lessor_city': 'Chennai',
            'lessor_pincode': '600001',
            'lessee_address': 'No. 456, Second Street',
            'lessee_city': 'Chennai',
            'lessee_pincode': '600002',
            'property_address': 'No. 789, Property Street',
            'property_city': 'Chennai',
            'property_pincode': '600073',
            'lease_period': '2',
            'start_date': '1st April 2024',
            'end_date': end_date_str,
            'lease_amount': '25,000',
            'lease_amount_words': 'Twenty Five Thousand',
            'rent_due_date': '1st',
            'security_deposit': '50,000',
            'security_deposit_words': 'Fifty Thousand',
            'notice_period': '3',
            'number_of_rooms': '3'
        })
    
    # Apply language-specific translations
    translations = {
        'en': {
            'Father Name': 'Father Name',
            'Mr. Witness One': 'Mr. Witness One',
            'Chennai': 'Chennai'
        },
        'hi': {
            'Father Name': 'अपने पिता का नाम',
            'Mr. Witness One': 'श्री विजय एक',
            'Chennai': 'चेन्नई'
        },
        'bn': {
            'Father Name': 'আমার পিতার নাম',
            'Mr. Witness One': 'শ্রী বিজয় এক',
            'Chennai': 'চেন্নাই'
        },
        'te': {
            'Father Name': 'నా పిల్లి పేరు',
            'Mr. Witness One': 'శ్రీ విజయ ఒక',
            'Chennai': 'చెన్నై'
        },
        'mr': {
            'Father Name': 'माझा वडील यांचा नाव',
            'Mr. Witness One': 'श्री విజయ एक',
            'Chennai': 'चेन्नई'
        },
        'ur': {
            'Father Name': 'میرے والد کا نام',
            'Mr. Witness One': 'شری ویجی ہے',
            'Chennai': 'چینనాی'
        },
        'gu': {
            'Father Name': 'માઝા પિતાનું નામ',
            'Mr. Witness One': 'શ્રી વિજય એક',
            'Chennai': 'ચેન્નઈ'
        },
        'kn': {
            'Father Name': 'ನನ್ನ ಪಿತಾನ ಹೆಸರು',
            'Mr. Witness One': 'ಶ್ರೀ ವಿಜಯ ಒಂದು',
            'Chennai': 'ಚೆನ್ನಾಯ'
        },
        'or': {
            'Father Name': 'ଆମଦ୍ବାରା ପିତାଙ୍କ ନାମ',
            'Mr. Witness One': 'ଶ୍ରୀ ବିଜଯ଼ ଏକ',
            'Chennai': 'ଚେନ୍ନାଇ'
        },
        'ta': {
            'Father Name': 'என் தந்தை பெயர்',
            'Mr. Witness One': 'சிறுவர் விஜய் ஒன்று',
            'Chennai': 'சென்னை'
        }
    }

    for key, value in defaults.items():
        if isinstance(value, str) and value in translations.get(language, {}):
            defaults[key] = translations[language][value]
        elif isinstance(value, str) and value.lower() in translations.get(language, {}):
            defaults[key] = translations[language][value.lower()]

    return defaults

documents = {
    'rental_agreement': {
        'templates': {
            'en': 'rental_agreement_template.txt',
            'hi': 'hi/rental_agreement_template.txt',
            'bn': 'bn/rental_agreement_template.txt',
            'te': 'te/rental_agreement_template.txt',
            'mr': 'mr/rental_agreement_template.txt',
            'ur': 'ur/rental_agreement_template.txt',
            'gu': 'gu/rental_agreement_template.txt',
            'kn': 'kn/rental_agreement_template.txt',
            'or': 'or/rental_agreement_template.txt',
            'ta': 'ta/rental_agreement_template.txt'
        },
        'fields': ['owner_name', 'owner_age', 'owner_father', 'owner_address', 'owner_city', 'owner_pincode',
                   'renter_name', 'renter_age', 'renter_father', 'renter_address', 'renter_city', 'renter_pincode',
                   'property_address', 'property_city', 'property_pincode', 'start_date', 'effective_date',
                   'duration', 'renewal_period', 'rent_amount', 'rent_amount_words', 'rent_due_date',
                   'rent_increase_percentage', 'security_deposit', 'security_deposit_words', 'notice_period',
                   'jurisdiction', 'witness1_name', 'witness1_address', 'witness2_name', 'witness2_address',
                   'execution_date', 'execution_place', 'date', 'month', 'year']
    },
    'land_sale_deed': {
        'templates': {
            'en': 'land_sale_deed_template.txt',
            'hi': 'hi/land_sale_deed_template.txt',
            'bn': 'bn/land_sale_deed_template.txt',
            'te': 'te/land_sale_deed_template.txt',
            'mr': 'mr/land_sale_deed_template.txt',
            'ur': 'ur/land_sale_deed_template.txt',
            'gu': 'gu/land_sale_deed_template.txt',
            'kn': 'kn/land_sale_deed_template.txt',
            'or': 'or/land_sale_deed_template.txt',
            'ta': 'ta/land_sale_deed_template.txt'
        },
        'fields': ['seller', 'seller_age', 'seller_father', 'seller_address', 'seller_city', 'seller_pincode',
                   'buyer', 'buyer_age', 'buyer_father', 'buyer_address', 'buyer_city', 'buyer_pincode',
                   'sale_amount', 'sale_amount_words', 'stamp_duty_bearer', 'registration_office',
                   'survey_number', 'area', 'north_boundary', 'south_boundary', 'east_boundary', 'west_boundary',
                   'property_address', 'property_city', 'property_pincode', 'witness1_name', 'witness1_address',
                   'witness2_name', 'witness2_address', 'execution_date', 'execution_place', 'date', 'month', 'year']
    },
    'power_of_attorney': {
        'templates': {
            'en': 'power_of_attorney_template.txt',
            'hi': 'hi/power_of_attorney_template.txt',
            'bn': 'bn/power_of_attorney_template.txt',
            'te': 'te/power_of_attorney_template.txt',
            'mr': 'mr/power_of_attorney_template.txt',
            'ur': 'ur/power_of_attorney_template.txt',
            'gu': 'gu/power_of_attorney_template.txt',
            'kn': 'kn/power_of_attorney_template.txt',
            'or': 'or/power_of_attorney_template.txt',
            'ta': 'ta/power_of_attorney_template.txt'
        },
        'fields': ['principal', 'principal_age', 'principal_father', 'principal_address', 'principal_city', 'principal_pincode',
                   'attorney', 'attorney_age', 'attorney_father', 'attorney_address', 'attorney_city', 'attorney_pincode',
                   'matter_description', 'effective_date', 'expiry_date', 'registration_office', 'stamp_duty_bearer',
                   'witness1_name', 'witness1_address', 'witness2_name', 'witness2_address', 'execution_date', 'execution_place',
                   'date', 'month', 'year']
    },
    'house_lease': {
        'templates': {
            'en': 'house_lease_template.txt',
            'hi': 'hi/house_lease_template.txt',
            'bn': 'bn/house_lease_template.txt',
            'te': 'te/house_lease_template.txt',
            'mr': 'mr/house_lease_template.txt',
            'ur': 'ur/house_lease_template.txt',
            'gu': 'gu/house_lease_template.txt',
            'kn': 'kn/house_lease_template.txt',
            'or': 'or/house_lease_template.txt',
            'ta': 'ta/house_lease_template.txt'
        },
        'fields': ['lessor', 'lessor_age', 'lessor_father', 'lessor_address', 'lessor_city', 'lessor_pincode',
                   'lessee', 'lessee_age', 'lessee_father', 'lessee_address', 'lessee_city', 'lessee_pincode',
                   'property_address', 'property_city', 'property_pincode', 'lease_period', 'start_date', 'end_date',
                   'lease_amount', 'lease_amount_words', 'rent_due_date', 'security_deposit', 'security_deposit_words',
                   'notice_period', 'jurisdiction', 'number_of_rooms', 'witness1_name', 'witness1_address',
                   'witness2_name', 'witness2_address', 'execution_date', 'execution_place', 'date', 'month', 'year']
    }
}

@document_bp.route('/document/<doc_type>')
def document_form(doc_type):
    if doc_type not in documents:
        return render_template('index.html', error='Invalid document type')
    language_mapping = {
        'en': 'English',
        'hi': 'Hindi',
        'bn': 'Bengali',
        'te': 'Telugu',
        'mr': 'Marathi',
        'ur': 'Urdu',
        'gu': 'Gujarati',
        'kn': 'Kannada',
        'or': 'Odia',
        'ta': 'Tamil'
    }
    languages = [{'code': lang, 'name': language_mapping.get(lang, lang.upper())} for lang in documents[doc_type]['templates'].keys()]
    return render_template('document_form.html', doc_type=doc_type, fields=documents[doc_type]['fields'], languages=languages)

@document_bp.route('/generate', methods=['POST'])
def generate_document():
    print("DEBUG: generate_document route called")
    doc_type = request.form.get('doc_type')
    language = request.form.get('language', 'en')
    if not doc_type or doc_type not in documents:
        return render_template('index.html', error='Invalid document type')

    # Collect field values and check for missing ones
    fields = documents[doc_type]['fields']
    data = {field: request.form.get(field, '').strip() for field in fields}

    # Map form fields to template placeholders
    field_mapping = {
        'rental_agreement': {
            'owner_name': 'landlord',
            'owner_age': 'landlord_age',
            'owner_father': 'landlord_father',
            'owner_address': 'landlord_address',
            'owner_city': 'landlord_city',
            'owner_pincode': 'landlord_pincode',
            'renter_name': 'tenant',
            'renter_age': 'tenant_age',
            'renter_father': 'tenant_father',
            'renter_address': 'tenant_address',
            'renter_city': 'tenant_city',
            'renter_pincode': 'tenant_pincode',
            'property_address': 'property_address',
            'property_city': 'property_city',
            'property_pincode': 'property_pincode',
            'start_date': 'start_date',
            'effective_date': 'effective_date',
            'duration': 'duration',
            'renewal_period': 'renewal_period',
            'rent_amount': 'rent_amount',
            'rent_amount_words': 'rent_amount_words',
            'rent_due_date': 'rent_due_date',
            'rent_increase_percentage': 'rent_increase_percentage',
            'security_deposit': 'security_deposit',
            'security_deposit_words': 'security_deposit_words',
            'notice_period': 'notice_period',
            'jurisdiction': 'jurisdiction',
            'witness1_name': 'witness1_name',
            'witness1_address': 'witness1_address',
            'witness2_name': 'witness2_name',
            'witness2_address': 'witness2_address',
            'execution_date': 'execution_date',
            'execution_place': 'execution_place',
            'date': 'date',
            'month': 'month',
            'year': 'year'
        },
        'land_sale_deed': {
            'seller': 'seller',
            'seller_age': 'seller_age',
            'seller_father': 'seller_father',
            'seller_address': 'seller_address',
            'seller_city': 'seller_city',
            'seller_pincode': 'seller_pincode',
            'buyer': 'buyer',
            'buyer_age': 'buyer_age',
            'buyer_father': 'buyer_father',
            'buyer_address': 'buyer_address',
            'buyer_city': 'buyer_city',
            'buyer_pincode': 'buyer_pincode',
            'sale_amount': 'sale_amount',
            'sale_amount_words': 'sale_amount_words',
            'stamp_duty_bearer': 'stamp_duty_bearer',
            'registration_office': 'registration_office',
            'survey_number': 'survey_number',
            'area': 'area',
            'north_boundary': 'north_boundary',
            'south_boundary': 'south_boundary',
            'east_boundary': 'east_boundary',
            'west_boundary': 'west_boundary',
            'property_address': 'property_address',
            'property_city': 'property_city',
            'property_pincode': 'property_pincode',
            'witness1_name': 'witness1_name',
            'witness1_address': 'witness1_address',
            'witness2_name': 'witness2_name',
            'witness2_address': 'witness2_address',
            'execution_date': 'execution_date',
            'execution_place': 'execution_place',
            'date': 'date',
            'month': 'month',
            'year': 'year'
        },
        'power_of_attorney': {
            'principal': 'principal',
            'principal_age': 'principal_age',
            'principal_father': 'principal_father',
            'principal_address': 'principal_address',
            'principal_city': 'principal_city',
            'principal_pincode': 'principal_pincode',
            'attorney': 'attorney',
            'attorney_age': 'attorney_age',
            'attorney_father': 'attorney_father',
            'attorney_address': 'attorney_address',
            'attorney_city': 'attorney_city',
            'attorney_pincode': 'attorney_pincode',
            'matter_description': 'matter_description',
            'effective_date': 'effective_date',
            'expiry_date': 'expiry_date',
            'registration_office': 'registration_office',
            'stamp_duty_bearer': 'stamp_duty_bearer',
            'witness1_name': 'witness1_name',
            'witness1_address': 'witness1_address',
            'witness2_name': 'witness2_name',
            'witness2_address': 'witness2_address',
            'execution_date': 'execution_date',
            'execution_place': 'execution_place',
            'date': 'date',
            'month': 'month',
            'year': 'year'
        },
        'house_lease': {
            'lessor': 'lessor',
            'lessor_age': 'lessor_age',
            'lessor_father': 'lessor_father',
            'lessor_address': 'lessor_address',
            'lessor_city': 'lessor_city',
            'lessor_pincode': 'lessor_pincode',
            'lessee': 'lessee',
            'lessee_age': 'lessee_age',
            'lessee_father': 'lessee_father',
            'lessee_address': 'lessee_address',
            'lessee_city': 'lessee_city',
            'lessee_pincode': 'lessee_pincode',
            'property_address': 'property_address',
            'property_city': 'property_city',
            'property_pincode': 'property_pincode',
            'lease_period': 'lease_period',
            'start_date': 'start_date',
            'end_date': 'end_date',
            'lease_amount': 'lease_amount',
            'lease_amount_words': 'lease_amount_words',
            'rent_due_date': 'rent_due_date',
            'security_deposit': 'security_deposit',
            'security_deposit_words': 'security_deposit_words',
            'notice_period': 'notice_period',
            'jurisdiction': 'jurisdiction',
            'number_of_rooms': 'number_of_rooms',
            'witness1_name': 'witness1_name',
            'witness1_address': 'witness1_address',
            'witness2_name': 'witness2_name',
            'witness2_address': 'witness2_address',
            'execution_date': 'execution_date',
            'execution_place': 'execution_place',
            'date': 'date',
            'month': 'month',
            'year': 'year'
        }
    }

    if doc_type in field_mapping:
        mapped_data = {}
        for form_field, template_field in field_mapping[doc_type].items():
            mapped_data[template_field] = data.get(form_field, '')
        data = mapped_data

    # Translate data if language is not English
    if language != 'en':
        print(f"Translating data to {language}")
        translated_data = {}
        for key, value in data.items():
            if value.strip():  # Only translate non-empty values
                translated = translate_text(value, language)
                translated_data[key] = translated
                print(f"Translated {key}: '{value}' -> '{translated}'")
            else:
                translated_data[key] = value
        data = translated_data
        print(f"Translation complete for {language}")

    # No longer require all fields to be filled. Missing fields will simply be empty in the template.
    # missing_fields = [field for field, value in data.items() if not value]

    # if missing_fields:
    #     error_msg = f"Please fill in all required fields: {', '.join(missing_fields)}"
    #     languages = list(documents[doc_type]['templates'].keys())
    #     return render_template('document_form.html', doc_type=doc_type, fields=fields, error=error_msg, values=data, languages=languages, selected_language=language)

    try:
        document = processor.generate_document(doc_type, data, language)
        # Use spaCy to extract entities
        doc_nlp = nlp(document)
        entities = [(ent.text, ent.label_) for ent in doc_nlp.ents]

        # Log history and save document
        print(f"DEBUG: Session contents: {dict(session)}")
        if 'user_id' in session:
            print(f"DEBUG: Found user_id in session: {session['user_id']}")
            
            # Add to history
            result = add_user_history(session['user_id'], 'generate_document', f'Generated {doc_type} in {language}')
            print(f"DEBUG: History add result: {result}")
            
            # Save generated document
            title = f"{doc_type.replace('_', ' ').title()} - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            doc_result = save_generated_document(
                session['user_id'], 
                doc_type, 
                language, 
                title, 
                document, 
                data  # Save the form data
            )
            print(f"DEBUG: Document save result: {doc_result}")
        else:
            print("DEBUG: No user_id found in session - user not logged in")

        return render_template('view_document.html', doc_type=doc_type, content=document, entities=entities)
    except Exception as e:
        import traceback
        error_message = f"Error generating document: {str(e)}"
        print(traceback.format_exc())
        flash(error_message, 'danger')
        languages = list(documents[doc_type]['templates'].keys())
        return render_template('document_form.html', doc_type=doc_type, fields=fields, error=error_message, values=data, languages=languages, selected_language=language)

@document_bp.route('/generate_from_prompt', methods=['POST'])
def generate_from_prompt():
    prompt = request.form.get('prompt', '').strip()
    if not prompt:
        flash('Please enter a prompt describing the document you want to generate.', 'error')
        return render_template('index.html')

    try:
        # Extract language from prompt, default to 'en'
        language_match = re.search(r'in\s+([a-zA-Z]+)', prompt)
        language = language_match.group(1).lower() if language_match else 'en'

        # Classify document type
        doc_type = processor.classify_document_type(prompt)
        if not doc_type:
            flash('Could not determine document type from your prompt. Please try rephrasing.', 'error')
            return render_template('index.html')

        # Extract entities
        entities = processor.extract_entities(prompt)

        # Generate document
        document = processor.generate_document(doc_type, entities, language=language)

        # Extract entities from generated document for display
        doc_nlp = nlp(document)
        extracted_entities = [(ent.text, ent.label_) for ent in doc_nlp.ents]

        # Log history and save document
        if 'user_id' in session:
            # Add to history
            add_user_history(session['user_id'], 'generate_from_prompt', f'Generated {doc_type} from prompt in {language}')
            
            # Save generated document
            title = f"{doc_type.replace('_', ' ').title()} from Prompt - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            save_generated_document(
                session['user_id'], 
                doc_type, 
                language, 
                title, 
                document, 
                entities  # Save the extracted entities
            )

        flash(f'Document type classified as: {doc_type.replace("_", " ").title()}', 'success')
        return render_template('view_document.html', doc_type=doc_type, content=document, entities=extracted_entities, prompt=prompt)

    except Exception as e:
        flash(f'Error generating document: {str(e)}', 'error')
        return render_template('index.html')

@document_bp.route('/download/<doc_type>/<format>')
def download_document(doc_type, format):
    # This would need the document content stored in session or database
    # For now, return a placeholder
    flash('Download functionality will be implemented with proper session management.', 'info')
    return render_template('index.html')

@document_bp.route('/api/process-prompt', methods=['POST'])
def api_process_prompt():
    """Process user prompt and extract information"""
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        
        if not prompt:
            return jsonify({'error': 'No prompt provided'}), 400
        
        # Classify document type
        doc_type = processor.classify_document_type(prompt)
        
        # Extract entities
        entities = processor.extract_entities(prompt)
        
        # Identify missing fields
        missing_fields = processor.identify_missing_fields(doc_type, entities)
        
        # Prepare response
        response = {
            'document_type': doc_type,
            'extracted_entities': entities,
            'missing_fields': missing_fields,
            'status': 'success'
        }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@document_bp.route('/api/generate-document', methods=['POST'])
def api_generate_document():
    """Generate final document with all required data"""
    try:
        data = request.get_json()
        doc_type = data.get('document_type')
        filled_data = data.get('filled_data', {})
        format_type = data.get('format', 'docx')
        language = data.get('language', 'en')
        custom_template = data.get('custom_template')
        
        if not doc_type or not filled_data:
            return jsonify({'error': 'Missing required data'}), 400
        
        # Generate document content
        if custom_template:
            from jinja2 import Template
            try:
                template = Template(custom_template)
                document_content = template.render(**filled_data)
            except Exception as e:
                return jsonify({'error': f'Error rendering custom template: {str(e)}'}), 400
        elif 'content' in filled_data:
            # Use provided content directly (for downloads from view_document.html)
            document_content = filled_data['content']
        else:
            # Merge filled_data with default values for a complete document
            complete_data = get_default_data_for_document(doc_type, language)
            complete_data.update(filled_data)  # User data overrides defaults
            document_content = processor.generate_document(doc_type, complete_data, language=language)
        
        # Log history and save document
        print(f"DEBUG API: Session contents: {dict(session)}")
        if 'user_id' in session:
            print(f"DEBUG API: Found user_id in session: {session['user_id']}")
            
            # Add to history
            add_user_history(session['user_id'], 'download_document', f'Downloaded {doc_type} as {format_type} in {language}')
            
            # Save generated document if not already saved
            if 'content' not in filled_data:  # Only save if it's a new generation
                title = f"{doc_type.replace('_', ' ').title()} - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                save_generated_document(
                    session['user_id'], 
                    doc_type, 
                    language, 
                    title, 
                    document_content, 
                    filled_data
                )
        else:
            print("DEBUG API: No user_id found in session - user not logged in")
        
        # Create file based on format
        if format_type == 'docx':
            return create_docx_file(document_content, doc_type)
        elif format_type == 'pdf':
            return create_pdf_file(document_content, doc_type)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def create_docx_file(content, doc_type):
    """Create DOCX file from content"""
    try:
        doc = Document()
        doc.add_heading(f'{doc_type.replace("_", " ").title()}', 0)
        
        # Add content
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        from flask import send_file
        from datetime import datetime
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'{doc_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': f'Error creating DOCX: {str(e)}'}), 500

def create_pdf_file(content, doc_type):
    """Create PDF file from content"""
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_path = temp_file.name
        temp_file.close()

        doc = SimpleDocTemplate(temp_path, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        body_style = ParagraphStyle(
            name='Body',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
        )
        title_style = ParagraphStyle(
            name='Title',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            spaceAfter=12,
        )

        story = []
        story.append(Paragraph(doc_type.replace('_', ' ').title(), title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Convert plain text line breaks to simple paragraphs
        for block in content.split('\n\n'):
            block_html = block.strip().replace('\n', '<br/>')
            if not block_html:
                continue
            story.append(Paragraph(block_html, body_style))
            story.append(Spacer(1, 0.12 * inch))

        doc.build(story)

        from flask import send_file
        from datetime import datetime
        return send_file(
            temp_path,
            as_attachment=True,
            download_name=f'{doc_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'error': f'Error creating PDF: {str(e)}'}), 500

@document_bp.route('/edit_document', methods=['GET', 'POST'])
def edit_document():
    doc_type = request.args.get('doc_type')
    content = request.args.get('content')

    if request.method == 'POST':
        edited_content = request.form.get('edited_content')
        doc_type = request.form.get('doc_type')
        return render_template('view_document.html', doc_type=doc_type, content=edited_content, entities=[]) # Simplified entities for now

    return render_template('edit_document.html', doc_type=doc_type, content=content)

@document_bp.route('/api/document/<doc_id>/view')
def api_view_document(doc_id):
    """View a saved document"""
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Authentication required'}), 401
        
        from app.models.history import supabase
        if not supabase:
            return jsonify({'error': 'Database connection failed'}), 500
        
        # Get document from database
        response = supabase.table('generated_documents').select('*').eq('id', doc_id).eq('user_id', session['user_id']).execute()
        
        if not response.data:
            return jsonify({'error': 'Document not found'}), 404
        
        document = response.data[0]
        
        # Return HTML view of the document
        return render_template('view_document.html', 
                             doc_type=document['document_type'], 
                             content=document['content'], 
                             entities=[],
                             title=document['title'])
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@document_bp.route('/api/document/<doc_id>/download/<format>')
def api_download_document(doc_id, format):
    """Download a saved document"""
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Authentication required'}), 401
        
        from app.models.history import supabase
        if not supabase:
            return jsonify({'error': 'Database connection failed'}), 500
        
        # Get document from database
        response = supabase.table('generated_documents').select('*').eq('id', doc_id).eq('user_id', session['user_id']).execute()
        
        if not response.data:
            return jsonify({'error': 'Document not found'}), 404
        
        document = response.data[0]
        
        # Log download activity
        add_user_history(session['user_id'], 'download_saved_document', f'Downloaded saved {document["document_type"]} as {format}')
        
        # Create file based on format
        if format == 'docx':
            return create_docx_file(document['content'], document['document_type'])
        elif format == 'pdf':
            return create_pdf_file(document['content'], document['document_type'])
        else:
            return jsonify({'error': 'Unsupported format'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
