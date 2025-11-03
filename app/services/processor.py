import os
from datetime import datetime
import re
import spacy
from jinja2 import Template
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# Load spaCy model for NLP processing
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

from app.services.document_generator import DocumentGenerator

class LegalDocumentProcessor:
    def __init__(self):
        self.document_types = {
            'rental_agreement': {
                'keywords': ['rental', 'rent', 'lease', 'tenant', 'landlord', 'monthly'],
                'required_fields': ['landlord', 'landlord_address', 'tenant', 'tenant_address', 'property_address', 'rent_amount', 'start_date', 'duration'],
                'template': 'rental_agreement_template.txt'
            },
            'land_sale_deed': {
                'keywords': ['sale', 'deed', 'property', 'buyer', 'seller', 'purchase'],
                'required_fields': ['seller', 'seller_address', 'buyer', 'buyer_address', 'property_address', 'sale_amount'],
                'template': 'land_sale_deed_template.txt'
            },
            'power_of_attorney': {
                'keywords': ['power', 'attorney', 'delegate', 'authority', 'behalf'],
                'required_fields': ['principal', 'principal_address', 'attorney', 'attorney_address', 'matter_description', 'effective_date', 'expiry_date'],
                'template': 'power_of_attorney_template.txt'
            },
            'house_lease': {
                'keywords': ['house', 'lease', 'lessor', 'lessee', 'property'],
                'required_fields': ['lessor', 'lessor_age', 'lessor_father', 'lessor_address', 'lessor_city', 'lessor_pincode', 'lessee', 'lessee_age', 'lessee_father', 'lessee_address', 'lessee_city', 'lessee_pincode', 'property_address', 'property_city', 'property_pincode', 'lease_period', 'start_date', 'end_date', 'lease_amount', 'lease_amount_words', 'rent_due_date', 'security_deposit', 'security_deposit_words', 'notice_period', 'number_of_rooms'],
                'template': 'house_lease_template.txt'
            }
        }
        self.document_generator = DocumentGenerator()

        self.entity_patterns = {
            'amounts': r'(?:Rs\.?|INR)?\s*(\d+(?:,\d+)*(?:\.\d{2})?)\s*(?:rupees?|Rs\.?|INR)?',
            'dates': r'\d{1,2}[-/]\d{1,2}[-/]\d{4}|\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}',
            'locations': r'\b(?:at|in)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            'durations': r'(\d+)\s*(?:year|month|week|day)s?'
        }

    def classify_document_type(self, prompt):
        """Classify the document type based on the user's prompt"""
        prompt_lower = prompt.lower()

        # Count keyword matches for each document type
        scores = {}
        for doc_type, info in self.document_types.items():
            score = sum(1 for keyword in info['keywords'] if keyword in prompt_lower)
            scores[doc_type] = score

        # Return the document type with the highest score
        if scores:
            return max(scores, key=scores.get)
        return None

    def extract_entities(self, prompt):
        """Extract entities from the user's prompt using spaCy and regex patterns"""
        doc = nlp(prompt)
        entities = {}

        # Extract named entities using spaCy
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                if 'landlord' not in entities:
                    entities['landlord'] = ent.text
                elif 'tenant' not in entities:
                    entities['tenant'] = ent.text
                elif 'seller' not in entities:
                    entities['seller'] = ent.text
                elif 'buyer' not in entities:
                    entities['buyer'] = ent.text
                elif 'principal' not in entities:
                    entities['principal'] = ent.text
                elif 'attorney' not in entities:
                    entities['attorney'] = ent.text
                elif 'lessor' not in entities:
                    entities['lessor'] = ent.text
                elif 'lessee' not in entities:
                    entities['lessee'] = ent.text
            elif ent.label_ == 'GPE' or ent.label_ == 'LOC':
                # Try to assign to specific address fields first
                if 'landlord_address' not in entities: # Added specific address field
                    entities['landlord_address'] = ent.text
                elif 'tenant_address' not in entities: # Added specific address field
                    entities['tenant_address'] = ent.text
                elif 'seller_address' not in entities: # Added specific address field
                    entities['seller_address'] = ent.text
                elif 'buyer_address' not in entities: # Added specific address field
                    entities['buyer_address'] = ent.text
                elif 'principal_address' not in entities: # Added specific address field
                    entities['principal_address'] = ent.text
                elif 'attorney_address' not in entities: # Added specific address field
                    entities['attorney_address'] = ent.text
                elif 'lessor_address' not in entities: # Added specific address field
                    entities['lessor_address'] = ent.text
                elif 'lessee_address' not in entities: # Added specific address field
                    entities['lessee_address'] = ent.text
                elif 'property_address' not in entities:
                    entities['property_address'] = ent.text
                elif 'address' not in entities:
                    entities['address'] = ent.text

        # Extract amounts using regex
        amount_match = re.search(self.entity_patterns['amounts'], prompt, re.IGNORECASE)
        if amount_match:
            entities['rent_amount'] = amount_match.group(1)
            entities['sale_amount'] = amount_match.group(1)
            entities['lease_amount'] = amount_match.group(1)

        # Extract dates
        date_match = re.search(self.entity_patterns['dates'], prompt, re.IGNORECASE)
        if date_match:
            entities['start_date'] = date_match.group(0)
            entities['effective_date'] = date_match.group(0)
            entities['expiry_date'] = date_match.group(0)
            entities['sale_date'] = date_match.group(0)

        # Extract durations
        duration_match = re.search(self.entity_patterns['durations'], prompt, re.IGNORECASE)
        if duration_match:
            entities['duration'] = duration_match.group(0)
            entities['renewal_period'] = duration_match.group(0)
            entities['lease_period'] = duration_match.group(0)
            entities['notice_period'] = duration_match.group(0)

        # Extract other specific fields if they are present in the prompt
        # For example, extracting property description for land_sale_deed
        if 'property_description' not in entities:
            # This is a simple placeholder. More sophisticated NLP might be needed for actual extraction.
            # For now, if the prompt contains a phrase like "a property located at X", we can try to extract X
            match = re.search(r'a property located at (.+?)(?:\.|,|$)', prompt, re.IGNORECASE)
            if match:
                entities['property_description'] = match.group(1).strip()
        
        if 'matter_description' not in entities:
            match = re.search(r'for (.+?) purposes', prompt, re.IGNORECASE)
            if match:
                entities['matter_description'] = match.group(1).strip() + ' purposes'
        
        # You might need more specific regex or NLP rules for other fields like father's name, city, pincode, etc.
        # For now, let's assume these would be provided through the form or default values.

        return entities

    def identify_missing_fields(self, doc_type, entities):
        """Identify missing required fields for the document type"""
        required_fields = self.document_types[doc_type]['required_fields']
        missing_fields = []
        
        for field in required_fields:
            if field not in entities or not entities[field]:
                missing_fields.append(field)
        
        return missing_fields

    def generate_document(self, doc_type, entities, language='en'):
        """Generate document content by filling the template with extracted entities"""
        if doc_type not in self.document_types:
            raise ValueError(f"Unsupported document type: {doc_type}")

        # Delegate to DocumentGenerator for multi-language support
        return self.document_generator.generate_document(doc_type, entities, language)

    def generate_docx(self, content, filename):
        """Generate a .docx file from the document content"""
        doc = Document()
        doc.add_heading('Legal Document', 0)

        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(filename)

    def generate_pdf(self, content, filename):
        """Generate a .pdf file from the document content"""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()

        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        story = []

        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))

        doc.build(story)
